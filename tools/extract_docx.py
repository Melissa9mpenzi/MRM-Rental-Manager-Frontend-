from __future__ import annotations

from pathlib import Path

from docx import Document


def extract_docx_to_text(src: Path) -> str:
    doc = Document(str(src))

    lines: list[str] = []

    # Paragraphs
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            lines.append(t)

    # Tables (simple pipe-delimited dump)
    for ti, table in enumerate(doc.tables, start=1):
        lines.append("")
        lines.append(f"[TABLE {ti}]")
        for row in table.rows:
            cells = [((c.text or "").replace("\n", " ").strip()) for c in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))

    return ("\n".join(lines).strip() + "\n") if lines else ""


def main() -> None:
    repo_root = Path(__file__).resolve().parents[1]
    src = repo_root / "rentalmgr_guide_v2.docx"
    out = repo_root / "rentalmgr_guide_v2.extracted.txt"

    if not src.exists():
        raise SystemExit(f"Input not found: {src}")

    text = extract_docx_to_text(src)
    out.write_text(text, encoding="utf-8")
    print(f"Wrote: {out} ({len(text)} chars)")


if __name__ == "__main__":
    main()

